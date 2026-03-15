from pathlib import Path
from langchain_core.documents import Document
import docx


def load_docs(path="./data/uploads"):
    """Load DOCX files using python-docx (no unstructured dependency)"""
    docs = []
    path_obj = Path(path)
    
    if not path_obj.exists():
        return docs
    
    for docx_file in path_obj.glob("**/*.docx"):
        try:
            doc = docx.Document(str(docx_file))
            text = "\n".join([para.text for para in doc.paragraphs])
            
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(docx_file), "type": "docx"}
                ))
        except Exception as e:
            print(f"Error loading {docx_file}: {e}")
    
    return docs
